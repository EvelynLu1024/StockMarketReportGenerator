import os
import pandas as pd
from datetime import datetime
from WindPy import w
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn


class MorningTalkWeekly:
    def __init__(self, start_date, end_date, yaml_path, output_dir):
        self.start_date = start_date
        self.end_date = end_date
        self.column_name = f"{start_date}_{end_date}"
        self.current_date = datetime.now().strftime("%Y年%m月%d日")
        self.current_date_str = datetime.now().strftime("%Y%m%d")
        self.date_difference = (datetime.strptime(end_date, "%Y%m%d") - datetime.strptime(start_date,
                                                                                          "%Y%m%d")).days + 1
        self.yaml_path = yaml_path
        self.output_dir = output_dir
        self.yaml_data = self.load_yaml()

        # Initialize WindPy
        w.start()

    def load_yaml(self):
        with open(self.yaml_path, 'r', encoding='utf-8') as file:
            return yaml.safe_load(file)

    def get_zdfweekly_w(self, code):
        zdf_w = w.wss(code, "pct_chg_per", f"startDate={self.start_date};endDate={self.end_date}")
        zdf = pd.DataFrame(zdf_w.Data, columns=zdf_w.Codes, index=[self.column_name]).T
        zdf = zdf / 100  # 去单位
        return zdf

    @staticmethod
    def sign_transformation(value):
        if value > 0:
            return f"涨{value:.2%}"
        elif value < 0:
            return f"跌{abs(value):.2%}"
        else:
            return "平收"

    def describe_indus(self, df):
        sorted_df = df.sort_values(by=self.column_name, ascending=False)

        up_top5 = sorted_df[sorted_df[self.column_name] > 0].head(5)
        down_top5 = sorted_df[sorted_df[self.column_name] < 0].tail(5).sort_values(by=self.column_name)

        if up_top5.empty:
            up_description = "没有行业上涨。"
        elif len(up_top5) < 5:
            up_description = f"仅{len(up_top5)}个行业上涨，" + "，".join(
                [f"{row['中文简称']}{row['涨跌幅']}" for _, row in up_top5.iterrows()]) + "。"
        else:
            up_description = "上涨前5位的行业分别是" + "，".join(
                [f"{row['中文简称']}{row['涨跌幅']}" for _, row in up_top5.iterrows()]) + "。"

        if down_top5.empty:
            down_description = "没有行业下跌。"
        elif len(down_top5) < 5:
            down_description = f"仅{len(down_top5)}个行业下跌，" + "，".join(
                [f"{row['中文简称']}{row['涨跌幅']}" for _, row in down_top5.iterrows()]) + "。"
        else:
            down_description = "下跌前5位的行业分别是" + "，".join(
                [f"{row['中文简称']}{row['涨跌幅']}" for _, row in down_top5.iterrows()]) + "。"

        return up_description + " " + down_description

    def describe_wind_index(self, df, top_n = 15):
        sorted_df = df.sort_values(by=self.column_name, ascending=False)

        up_top = sorted_df[sorted_df[self.column_name] > 0].head(top_n)
        down_top = sorted_df[sorted_df[self.column_name] < 0].tail(top_n).sort_values(by=self.column_name)

        if up_top.empty:
            up_description = "全部下跌。"
        elif len(up_top) < top_n:
            up_description = f"仅{len(up_top)}个概念上涨，" + "，".join(
                [f"{row['中文简称']}{row['涨跌幅']}" for _, row in up_top.iterrows()]) + "。"
        else:
            up_description = f"上涨前{top_n}位的概念分别是" + "，".join(
                [f"{row['中文简称']}{row['涨跌幅']}" for _, row in up_top.iterrows()]) + "。"

        if down_top.empty:
            down_description = "全部上涨。"
        elif len(down_top) < top_n:
            down_description = f"仅{len(down_top)}个概念下跌，" + "，".join(
                [f"{row['中文简称']}{row['涨跌幅']}" for _, row in down_top.iterrows()]) + "。"
        else:
            down_description = f"下跌前{top_n}位的概念分别是" + "，".join(
                [f"{row['中文简称']}{row['涨跌幅']}" for _, row in down_top.iterrows()]) + "。"

        return "wind热门概念"+ up_description + "\n   " + down_description

    def generate_word_report(self):
        doc = Document()

        # 添加标题
        title = doc.add_heading('', level=1)
        run = title.add_run('__权益投资部晨会纪要')
        run.font.name = '楷体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加日期
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(self.current_date)
        date_run.font.name = '楷体'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        date_run.font.size = Pt(12)
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 添加空行
        doc.add_paragraph()

        # 添加“市场概况”
        market_paragraph = doc.add_paragraph()
        market_run = market_paragraph.add_run('市场概况')
        market_run.font.bold = True
        market_run.font.size = Pt(14)
        market_run.font.name = '楷体'
        market_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        market_paragraph.paragraph_format.line_spacing = 1.5

        # 添加段落内容
        paragraphs = self.get_paragraphs()
        for text in paragraphs:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = '楷体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            run.font.size = Pt(12)
            paragraph.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
            paragraph.paragraph_format.line_spacing = 1.5  # 设置1.5倍行距
            if any(kw in text for kw in ["1、A股市场", "2、港股市场", "3、美股市场"]):  # 特定关键字加粗
                run.font.bold = True
            paragraph.paragraph_format.space_after = Pt(0)  # 设置段后间距

        # 保存文档
        doc.save(os.path.join(self.output_dir, f'某险资权益投资部晨会纪要（{self.current_date_str}）.docx'))

    def get_paragraphs(self):
        a_pct_chg = self.get_zdfweekly_w(','.join(list(self.yaml_data["a_index_codes"].values())))
        hk_pct_chg = self.get_zdfweekly_w(
            ','.join([code for category in self.yaml_data['hk_codes'].values() for code in category.values()]))

        szzz_zdf = self.sign_transformation(a_pct_chg.loc['000001.SH', self.column_name])
        szcz_zdf = self.sign_transformation(a_pct_chg.loc['399001.SZ', self.column_name])
        cybz_zdf = self.sign_transformation(a_pct_chg.loc['399006.SZ', self.column_name])
        kc50_zdf = self.sign_transformation(a_pct_chg.loc['000688.SH', self.column_name])
        zz500_zdf = self.sign_transformation(a_pct_chg.loc['000905.SH', self.column_name])
        zz1000_zdf = self.sign_transformation(a_pct_chg.loc['000852.SH', self.column_name])
        zz2000_zdf = self.sign_transformation(a_pct_chg.loc['932000.CSI', self.column_name])

        sszz_spj = w.wss(self.yaml_data["a_index_codes"]["上证指数"], "close",
                         f"tradeDate={self.end_date};priceAdj=U;cycle=D").Data[0][0]
        a_rjcje = w.wsee("a001010100000000", "sec_pq_amt_sum",
                         f"unit=1;startDate={self.start_date};endDate={self.end_date};currencyType=Cur=CNY;DynamicTime=1").Data[
                      0][0] / 100000000 / self.date_difference
        bxzj = w.wsee("1000025141000000", "sec_pq_ncashinflow_sum_chn",
                      f"startDate={self.start_date};endDate={self.end_date};DynamicTime=1").Data[0][
                   0] / 100000000 / self.date_difference

        sw_codes = ','.join(list(self.yaml_data["申万一级行业"].values()))
        sw_indus_pct_chg = self.get_zdfweekly_w(sw_codes)
        indus_info = {indus_code: indus_name for indus_name, indus_code in self.yaml_data['申万一级行业'].items()}
        sw_indus_pct_chg["涨跌幅"] = sw_indus_pct_chg[self.column_name].apply(lambda x: self.sign_transformation(x))
        sw_indus_pct_chg["中文简称"] = sw_indus_pct_chg.index.map(lambda x: indus_info[x])

        morning_talk_weekly_a = (
            f"上周（{self.start_date}-{self.end_date}），A股三大股指__。截至收盘，沪指{szzz_zdf}，报{sszz_spj:.2f}点，深证城指{szcz_zdf}，创业板指{cybz_zdf}，"
            f"科创50{kc50_zdf}，中证500{zz500_zdf}，中证1000{zz1000_zdf}，中证2000{zz2000_zdf}。"
            f"市场成交额日均成交额{a_rjcje:.2f}亿元，北向资金{bxzj:.2f}亿元。")

        wind_index_codes = ','.join(list(self.yaml_data["wind_index"].values()))
        sw_wind_index_pct_chg = self.get_zdfweekly_w(wind_index_codes)
        wind_index_info = {wind_index_code: wind_index_name for wind_index_name, wind_index_code in
                           self.yaml_data['wind_index'].items()}
        sw_wind_index_pct_chg["涨跌幅"] = sw_wind_index_pct_chg[self.column_name].apply(
            lambda x: self.sign_transformation(x))
        sw_wind_index_pct_chg["中文简称"] = sw_wind_index_pct_chg.index.map(lambda x: wind_index_info[x])

        morning_talk_weekly_aindus = self.describe_indus(sw_indus_pct_chg)
        morning_talk_weekly_wind_index = self.describe_wind_index(sw_wind_index_pct_chg)

        # 港股市场
        hszs_zdf = self.sign_transformation(a_pct_chg.loc['HSI.HI', self.column_name])
        hskjzs_zdf = self.sign_transformation(a_pct_chg.loc['HSTECH.HI', self.column_name])
        hsgqzs_zdf = self.sign_transformation(a_pct_chg.loc['HSCEI.HI', self.column_name])

        hszs_spj = w.wss(self.yaml_data["a_index_codes"]["恒生指数"], "close",
                         f"tradeDate={self.end_date};priceAdj=U;cycle=D").Data[0][0]
        nxzj = w.wsee("1000011937000000", "sec_pq_ncashinflow_sum_chn",
                      f"startDate={self.start_date};endDate={self.end_date};DynamicTime=1").Data[0][
                   0] / 100000000 / self.date_difference

        code_info = {code: {'second_layer': second_layer, 'chinese_name': company}
                     for second_layer, companies in self.yaml_data['hk_codes'].items()
                     for company, code in companies.items()}
        hk_pct_chg["indus_type"] = hk_pct_chg.index.map(lambda x: code_info[x]['second_layer'])
        hk_pct_chg["股票简称"] = hk_pct_chg.index.map(lambda x: code_info[x]['chinese_name'])
        hk_pct_chg = hk_pct_chg.sort_values(by=['indus_type', self.column_name], ascending=[False, False])
        hk_pct_chg["涨跌幅"] = hk_pct_chg[self.column_name].apply(lambda x: self.sign_transformation(x))

        morning_talk_weekly_h = (
            f"上周（{self.start_date}-{self.end_date}），港股三大股指__。截至收盘，恒生指数{hszs_zdf}，报{hszs_spj}点，"
            f"恒生科技指数{hskjzs_zdf}，恒生国企指数{hsgqzs_zdf}。南向资金{nxzj:.2f}亿港元。")

        # 生成港股文字描述
        description = {
            '港股科技股': [],
            '港股医药股': [],
            '港股内房股': []
        }
        for indus_type, group in hk_pct_chg.groupby('indus_type'):
            indus_description = f"港股{indus_type[2:]}：" + "，".join(
                [f"{row['股票简称'] + row['涨跌幅']}" for _, row in group.iterrows()])
            description[indus_type] = indus_description

        morning_talk_weekly_h_part2 = "；\n   ".join(description.values()) + "。"

        # 美股市场
        usa_codes = ','.join(list(self.yaml_data["usa_index_codes"].values()))
        usa_pct_chg = self.get_zdfweekly_w(usa_codes)
        dqs_zdf = self.sign_transformation(usa_pct_chg.loc['DJI.GI', self.column_name])
        bp500_zdf = self.sign_transformation(usa_pct_chg.loc['SPX.GI', self.column_name])
        nsdk_zdf = self.sign_transformation(usa_pct_chg.loc['IXIC.GI', self.column_name])

        morning_talk_weekly_usa = (f"上周（{self.start_date}-{self.end_date}），美股三大股指__。截至收盘，道指{dqs_zdf}，"
                                   f"标普500指数{bp500_zdf}，纳指{nsdk_zdf}。")

        # 推导式创建一个字典来存储代码和它们的中文简称以及第二层结构
        usa_code_info = {code: {'second_layer': second_layer, 'chinese_name': company}
                         for second_layer, companies in self.yaml_data['usa_codes'].items()
                         for company, code in companies.items()}
        # 提取usa_stocks代码
        usa_stock_codes = [code for category in self.yaml_data['usa_codes'].values() for code in category.values()]
        usa_stocks = self.get_zdfweekly_w(usa_stock_codes)
        usa_stocks["indus_type"] = usa_stocks.index.map(lambda x: usa_code_info[x]['second_layer'])
        usa_stocks["股票简称"] = usa_stocks.index.map(lambda x: usa_code_info[x]['chinese_name'])
        usa_stocks = usa_stocks.sort_values(by=['indus_type', self.column_name], ascending=[False, False])
        usa_stocks["涨跌幅"] = usa_stocks[self.column_name].apply(lambda x: self.sign_transformation(x))

        # 生成美股文字描述
        description = {
            '美股科技股': [],
            '美股中概股': []
        }
        for indus_type, group in usa_stocks.groupby('indus_type'):
            indus_description = f"{indus_type}：" + "，".join(
                [f"{row['股票简称'] + row['涨跌幅']}" for _, row in group.iterrows()])
            description[indus_type] = indus_description

        morning_talk_weekly_usa_part2 = "；\n   ".join(description.values()) + "。"

        # 返回段落内容
        return [
            "1、A股市场", morning_talk_weekly_a, morning_talk_weekly_aindus, morning_talk_weekly_wind_index,
            "2、港股市场", morning_talk_weekly_h, morning_talk_weekly_h_part2,
            "3、美股市场", morning_talk_weekly_usa, morning_talk_weekly_usa_part2,
            "（汇报人：）"
        ]


# if __name__ == "__main__":
#     # 设置工作目录
#     os.chdir("")
#
#     # 实例化MorningTalkWeekly类并生成报告
#     morningtalk_weekly = MorningTalkWeekly(
#         start_date="20240603",
#         end_date="20240607",
#         yaml_path='mt_weekly.yaml',
#         output_dir='.'
#     )
#     morningtalk_weekly.generate_word_report()
